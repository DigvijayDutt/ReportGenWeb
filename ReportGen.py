import re
import os
import sys
import shutil
import subprocess
import zipfile
import asyncio
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from nicegui import ui,app
from docx2pdf import convert
# -------------------- STYLE MAP --------------------
STYLE_MAP = {
    "Normal": {
        "font_name": "Times New Roman",
        "font_size": Pt(12),
        "bold": False,
        "italic": False,
        "underline": False,
        "color": RGBColor(0, 0, 0),
        "line_spacing": 1.15
    },
    "Heading 1": {
        "font_name": "Times New Roman",
        "font_size": Pt(18),
        "bold": True,
        "italic": False,
        "underline": True,
        "color": RGBColor(47, 84, 150),
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "line_spacing": 1.5
    },
    "Heading 3": {
        "font_name": "Times New Roman",
        "font_size": Pt(12),
        "bold": True,
        "italic": False,
        "underline": False,
        "color": RGBColor(47, 84, 150),
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "line_spacing": 1.5
    },
    "List Paragraph": {
        "font_name": "Times New Roman",
        "font_size": Pt(12),
        "bold": True,
        "italic": False,
        "underline": False,
        "color": RGBColor(0, 0, 0),
        "line_spacing": 1.15
    }
}


def apply_style(para, style_key="Normal"):
    """Apply a style from STYLE_MAP to a paragraph and its runs."""
    style = STYLE_MAP.get(style_key, STYLE_MAP["Normal"])
    para_format = para.paragraph_format
    para_format.line_spacing = style["line_spacing"]
    if "alignment" in style:
        para.alignment = style["alignment"]
    for run in para.runs:
        run.font.name = style["font_name"]
        run.font.size = style["font_size"]
        run.font.bold = style["bold"]
        run.font.italic = style["italic"]
        run.font.underline = style["underline"]
        run.font.color.rgb = style["color"]


def set_cell_padding(cell, top=200, start=200, bottom=200, end=200):
    """
    Set cell padding (in twips: 1/20th of a point).
    Example: 200 twips ≈ 10 pt ≈ 3.5 mm
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


# -------------------- FOLDER MANAGEMENT --------------------
def get_appdata_base(app_name="DocumentGenerator"):
    """Return a writable base folder for uploads/outputs, cross-platform."""
    try:
        if os.name == "nt":
            base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
        else:
            base = os.environ.get("XDG_DATA_HOME") or os.path.expanduser("~/.local/share")
        app_folder = os.path.join(base, app_name)
        os.makedirs(app_folder, exist_ok=True)
        return app_folder
    except Exception:
        import tempfile
        fallback = os.path.join(tempfile.gettempdir(), app_name)
        os.makedirs(fallback, exist_ok=True)
        return fallback


APPDATA_BASE = get_appdata_base()
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "outputs")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def resource_path(relative_path):
    try:
        if getattr(sys, "frozen", False):
            base_path = sys._MEIPASS
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


def clean_folder(folder_path):
    """Remove all files/subfolders inside a folder."""
    if not os.path.isdir(folder_path):
        return
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f"[WARNING] Failed to delete {file_path}. Reason: {e}")


# -------------------- DOCUMENT GENERATOR --------------------
def generate_docs(excel_path, room_image_map, output_filename, row_index, log_callback):
    try:
        df = pd.read_excel(excel_path)
        for col in df.select_dtypes(include=["datetime"]):
            df[col] = df[col].dt.strftime("%d-%m-%Y")
    except Exception as e:
        log_callback(f"[ERROR] Failed to read Excel file: {e}")
        return None

    try:
        columns = df.columns
        data = df.values
        if len(data) == 0:
            log_callback("[ERROR] Excel file contains no rows.")
            return None
        if row_index >= len(data):
            log_callback(f"[WARNING] Row index {row_index} out of bounds. Using last available row.")
            row_index = len(data) - 1

        template_path = resource_path(os.path.join("templates", "template.docx"))
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            raise FileNotFoundError("template.docx not found")
    except Exception:
        log_callback("[WARNING] Template missing or inaccessible. Using blank document.")
        doc = Document()

    # Heading
    heading = doc.add_heading("FIRST INSPECTION REPORT", level=1)
    apply_style(heading, "Heading 1")

    # Table
    table = doc.add_table(rows=4, cols=1)
    table.autofit = True
    nestedt = table.cell(0,0).add_table(rows=1,cols=2)

    # Header Image
    try:
        imageCell = table.cell(1, 0)
        imageCell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        home_tuple = next(((k, v) for k, v in room_image_map.items() if k.lower() == "home"), None)
        if home_tuple:
            para1 = imageCell.paragraphs[0]
            run = para1.add_run()
            run.add_break()
            run.add_break()
            try:
                run.add_picture(home_tuple[1][0])
                para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.add_break()
            except Exception as e:
                imageCell.text = "[Header image missing]"
                log_callback(f"[WARNING] Failed to insert header image: {e}")
        else:
            imageCell.text = "[Header image not provided]"
    except Exception as e:
        log_callback(f"[WARNING] Failed to load image header: {e}")

    note = ''
    # Fill metadata
    for i in range(len(columns)):
        col_name = re.sub(r"\s*\((?:select.*|dd/mm/yyyy)\)\s*:?", "", str(columns[i]), flags=re.IGNORECASE).lower().strip()
        display_col_name = re.sub(r"\s*\((?:select.*|dd/mm/yyyy)\)\s*:?", "", str(columns[i]), flags=re.IGNORECASE).strip()
        try:
            if col_name in ["policyholder", "address", "insurer", "adjuster","description of risk","claim #","date of report",
                            "date assigned","date of inspection", "date of loss", "type of loss", "cause of loss","assigned gc","pm contact"]:
                textCell = nestedt.cell(0,0)
                if col_name == "cause of loss":
                    textCell = table.cell(3, 0)
                if col_name == "description of risk":
                    textCell = table.cell(2, 0)
                if col_name in ["claim #","date of report","date assigned","assigned gc","pm contact"]:
                    textCell = nestedt.cell(0,1)
                para = textCell.add_paragraph()
                run = para.add_run(f"{display_col_name}: ")
                if col_name == "cause of loss" or col_name == "description of risk":
                    para.add_run(f"\n{str(data[row_index][i])}\n")
                else:
                    para.add_run(str(data[row_index][i]))
                apply_style(para, "Normal")
                run.bold = True
                if col_name == "cause of loss" or col_name == "description of risk":
                    run.font.color.rgb = RGBColor(47, 84, 150)
            elif col_name in ["indemnity reserves:", "expense reserves:"]:
                para = doc.add_paragraph(display_col_name + " ", style='List Bullet')
                if col_name == "indemnity reserves:":
                    temp = data[row_index][i].split("HST")
                    note = temp[1]
                    run = para.add_run(temp[0] + "HST")
                    apply_style(para, "List Paragraph")
                    run.font.bold = False
                else:
                    run = para.add_run(str(data[row_index][i]))
                    run1 = para.add_run(f"\nNote: {note.strip()}")
                    run1.add_break()
                    apply_style(para, "List Paragraph")
                    run.font.bold = False
                    run1.font.bold = False
            elif col_name == "product manager":
                temp = data[row_index][i].split('\n')

                para1 = nestedt.cell(0,0).add_paragraph()
                run2 = para1.add_run(f"{display_col_name}: ")
                para1.add_run(f"{temp[0]}")
                apply_style(para1, "Normal")
                run2.bold =  True

                para = doc.add_paragraph()
                para.add_run().add_break()
                run = para.add_run("Thank you,")
                run1 = para.add_run(f"\n{temp[0]}")
                for line in temp[1:]:
                    para.add_run(f"\n{line}")
                apply_style(para, "Normal")
                run1.font.bold = True
            else:
                if col_name == "conclusion":
                    doc.add_paragraph().add_run().add_break()
                    heading2 = doc.add_heading(display_col_name, level=2)
                    par = doc.add_paragraph(str(data[row_index][i]))
                    apply_style(heading2, "Heading 3")
                    apply_style(par, "Normal")
                else:
                    hed = doc.add_heading(display_col_name, level=2)
                    apply_style(hed, "Heading 3")
                    para = doc.add_paragraph(str(data[row_index][i]))
                    apply_style(para, "Normal")
                    if col_name != "recommended reserves for trinity's involvement:":
                        para.add_run().add_break()
                    else:
                        para.paragraph_format.line_spacing = 1.5
                
        except Exception as e:
            log_callback(f"[WARNING] Failed to insert column {display_col_name}: {e}")

    # Photos Section
    if any(room_image_map.values()):
        doc.add_page_break()
        phead = doc.add_heading("PHOTOGRAPHS", level=1)
        phead.add_run().add_break()
        apply_style(phead, "Heading 1")
        room_image_map = {k: v for k, v in room_image_map.items() if k.lower() != "home"}
        if not room_image_map:
            doc.add_paragraph("[No photographs provided]")

        for idx, (room_name, image_list) in enumerate(room_image_map.items()):
            rhead = doc.add_heading(room_name.upper(), level=3)
            rhead.add_run().add_break()
            apply_style(rhead, "Normal")
            rhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in rhead.runs:
                run.font.size = Pt(18)

            for i in range(0, len(image_list), 4):
                table = doc.add_table(rows=2, cols=2)
                table.autofit = True
                batch = image_list[i:i + 4]
                for row in range(2):
                    for col in range(2):
                        img_idx = row * 2 + col
                        if img_idx < len(batch):
                            cell = table.cell(row, col)
                            p = cell.paragraphs[0]
                            if col == 0:
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run()
                            set_cell_padding(cell)
                            try:
                                run.add_picture(batch[img_idx], width=Inches(2.5), height=Inches(2.5))
                            except Exception as e:
                                cell.text = "[Image missing]"
                                log_callback(f"[WARNING] Failed to insert image {batch[img_idx]}: {e}")
            if idx < len(room_image_map) - 1:
                doc.add_page_break()

    # Save document
    output_path = output_filename
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    log_callback(f"[SUCCESS] Document saved: {output_filename}")
    return output_filename

class MultiDocApp:
    def __init__(self):
        self.excel_path = None
        self.claim_image_map = {}
        self.output_folder = OUTPUT_FOLDER
        self.download_path = ''
        ui.colors(secondary='#033452')
        with ui.row().classes('w-full justify-center items-center gap-3'):
            ui.image('./logo.jpg').classes('!absolute left-4 top-3 w-35 h-15')
            ui.label('Report Automation Tool').classes('my-3 text-2xl font-bold text-black-700 mb-4 text-outline-black')

        with ui.column().classes('gap-4 w-2/3 mx-auto px-12 py-8 rounded-xl shadow-2xl bg-white'):
            self.render_widgets()

            ui.button('⚙️ Generate Documents', on_click=self.process_files).classes('bg-[#033452] text-white w-full').props('color=secondary')
            ui.button('⬇️ Download', on_click=self.get_download).classes('!bg-[#033452] text-white w-full')
            self.log_area = ui.html('', sanitize=False) \
                .classes('w-full h-64 bg-black rounded p-2 overflow-y-scroll')
            
    @ui.refreshable_method
    def render_widgets(self):  
        with ui.row().classes('gap-6 w-full justify-between'):
            with ui.column().classes('items-start'):
                ui.label('Upload an Excel file to generate multiple Word reports:')
                ui.upload(on_upload=self.handle_excel_upload, auto_upload=True).props('accept=.xls,.xlsx flat bordered color=secondary')

            with ui.column().classes('items-start'):
                ui.label('Upload image ZIP containing claim folders:')
                ui.upload(on_upload=self.handle_folder_zip_upload, auto_upload=True).props('accept=.zip flat bordered color=secondary')

    def log(self, message: str):
        color = "white"
        if message.startswith("[ERROR]"):
            color = "red"
        elif message.startswith("[SUCCESS]"):
            color = "lightgreen"
        elif message.startswith("[WARNING]"):
            color = "yellow"

        self.log_area.content += f'<span style="color:{color}">{message}</span><br>'
        self.log_area.update()


    def clear_log(self):
        # implement your own log clearing method; example:
        if hasattr(self, 'log_area'):
            self.log_area.content = ''
            self.log_area.update()

    async def handle_excel_upload(self, e):
        """Save uploaded Excel file (NiceGUI v2.2+ safe version)."""
        if not e.file:
            return

        # take the first uploaded file
        file = e.file

        # build save path
        path = os.path.join(UPLOAD_FOLDER, file.name)

        # save it to disk
        await file.save(path)

        self.excel_path = path
        self.log(f"[INFO] Excel file uploaded: {path}")
    
    async def handle_folder_zip_upload(self, e):
        """Accepts a ZIP of base_folder/parent_folder/ClaimA/Room1/... and builds claim_image_map."""
        try:
            if not getattr(e, 'file', None):
                self.log("[ERROR] No ZIP file received.")
                return

            file = e.file
            zip_name = file.name
            zip_path = os.path.join(UPLOAD_FOLDER, zip_name)
            await file.save(zip_path)

            self.log(f"[INFO] Received folder ZIP: {zip_name}")

            # Prepare extraction path
            extract_base = os.path.join(UPLOAD_FOLDER, os.path.splitext(zip_name)[0])
            if os.path.exists(extract_base):
                shutil.rmtree(extract_base)
            os.makedirs(extract_base, exist_ok=True)

            # Extract ZIP contents
            with zipfile.ZipFile(zip_path, 'r') as z:
                z.extractall(extract_base)

            # Identify parent folder automatically (first-level dir inside base)
            extracted_items = [d for d in os.listdir(extract_base) if os.path.isdir(os.path.join(extract_base, d))]
            if not extracted_items:
                self.log("[ERROR] ZIP did not contain any folders.")
                return
            parent_folder = os.path.join(extract_base, extracted_items[0])
            self.log(f"[INFO] Detected parent folder: {parent_folder}")

            # Build claim → room → images mapping
            valid_ext = ('.jpg', '.jpeg', '.png')
            self.claim_image_map = {}

            for claim in os.listdir(parent_folder):
                claim_path = os.path.join(parent_folder, claim)
                if not os.path.isdir(claim_path):
                    continue
                room_images = {}
                for room in os.listdir(claim_path):
                    room_path = os.path.join(claim_path, room)
                    if not os.path.isdir(room_path):
                        continue
                    images = sorted([
                        os.path.join(room_path, f)
                        for f in os.listdir(room_path)
                        if f.lower().endswith(valid_ext)
                    ])
                    if images:
                        room_images[room] = images
                if room_images:
                    self.claim_image_map[claim] = room_images

            if self.claim_image_map:
                self.log(f"[SUCCESS] Found {len(self.claim_image_map)} claim(s) inside {parent_folder}")
            else:
                self.log("[WARNING] No valid claim folders or images found.")

        except Exception as ex:
            self.log(f"[ERROR] Failed to process uploaded folder ZIP: {ex}")


    async def process_files(self):
        if not getattr(self, 'excel_path', None) or not getattr(self, 'claim_image_map', None) or not self.claim_image_map:
            self.log("[ERROR] Please upload an Excel file and a valid image folder structure (ZIP).")
            return

        try:
            # clear logs and optionally show status
            self.clear_log()
            self.log("[INFO] Preparing to generate documents...")

            # make sure upload/output folders exist
            os.makedirs(UPLOAD_FOLDER, exist_ok=True)
            os.makedirs(self.output_folder, exist_ok=True)

            generated_files = []

            # Iterate claims and generate docs
            for i, (claim_name, room_images) in enumerate(self.claim_image_map.items()):
                # sanitize claim_name to a safe filename
                safe_name = "".join(c if c.isalnum() or c in " ._-" else "_" for c in claim_name).strip()
                if not safe_name:
                    safe_name = f"claim_{i+1}"

                output_filename = os.path.join(self.output_folder, f"{safe_name}.docx")
                self.log(f"[INFO] Generating document for claim: {claim_name} -> {output_filename}")

                # call your document generator (synchronous). It uses `self.log` as callback.
                result = generate_docs(self.excel_path, room_images, output_filename, i, self.log)

                if result:
                    generated_files.append(result)
                else:
                    self.log(f"[WARNING] Document generation returned no result for claim: {claim_name}")

            # Post-processing: report and offer download
            if generated_files:
                self.log(f"[SUCCESS] {len(generated_files)} document(s) generated in {self.output_folder}")

                # create zip of generated files for single-download convenience
                zip_name = "Generated_Reports.zip"
                zip_path = os.path.join(self.output_folder, zip_name)
                # overwrite existing zip if present
                if os.path.exists(zip_path):
                    os.remove(zip_path)

                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for fpath in generated_files:
                        if os.path.exists(fpath):
                            zf.write(fpath, arcname=os.path.basename(fpath))
                self.log(f"[INFO] Packaged {len(generated_files)} documents into: {zip_path}")

                # opdoc = generated_files[0]
                # oppdf = opdoc.replace('.docx','.pdf')
                # await asyncio.to_thread(convert, opdoc,oppdf)
                # app.add_static_files('/outputs', 'outputs')
                # with ui.dialog() as preview_dialog, ui.card().classes('w-3/4 h-[80vh] p-4'):
                #     ui.label('📄 PDF Preview').classes('text-lg font-bold mb-2')
                #     ui.element('iframe').props('src="/outputs/claim1.pdf" type="application/pdf"').classes('w-full h-[70vh] rounded shadow-lg border')
                #     ui.button('Close', on_click=preview_dialog.close).props('color=primary text-white')
                # preview_dialog.open()


                self.download_path = zip_path
                
            else:
                self.log("[WARNING] No documents were generated.")

        except Exception as e:
            self.log(f"[ERROR] Unexpected error while processing files: {e}")
    
    
    async def get_download(self):
        if not self.download_path or not os.path.exists(self.download_path):
            self.log("[ERROR] Nothing to download.")
            return

        ui.download(self.download_path)
        self.log(f"[INFO] Download initiated: {self.download_path}")

        # async refresh — wait to ensure download starts before cleanup
        asyncio.create_task(self._refresh_after_download())

    async def _refresh_after_download(self):
        """Wait a short time, then reset uploads and clean folders."""
        await asyncio.sleep(2)  # give browser time to start download
        self.log("[INFO] Resetting upload widgets for next use...")
        self.excel_path = None
        self.claim_image_map = {}
        self.render_widgets.refresh()  # 🔄 refresh upload UI
        await asyncio.sleep(1)
        self.log("[INFO] Cleaning temporary folders...")
        clean_folder(UPLOAD_FOLDER)
        clean_folder(OUTPUT_FOLDER)
        self.clear_log()
        self.log("[SUCCESS] Ready for new upload cycle.")

# Run app
MultiDocApp()
# port = int(os.environ.get("PORT", 8080))
ui.run(host="127.0.0.1", port=9080, reload=False, workers=1)

